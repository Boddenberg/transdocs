from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DESTINO = (
    Path(__file__).resolve().parents[1]
    / "app"
    / "recursos"
    / "modelos"
    / "escritura_venda_compra_estruturada.docx"
)

SECOES = [
    (None, "[PREENCHER:ABERTURA_DO_ATO]"),
    ("DAS PARTES VENDEDORAS", "[PREENCHER:QUALIFICACAO_VENDEDORES]"),
    ("DAS PARTES COMPRADORAS", "[PREENCHER:QUALIFICACAO_COMPRADORES]"),
    ("DO IMÓVEL", "[PREENCHER:DESCRICAO_DO_IMOVEL]"),
    ("DO TÍTULO AQUISITIVO", "[PREENCHER:TITULO_AQUISITIVO]"),
    ("DO PREÇO E DA FORMA DE PAGAMENTO", "[PREENCHER:PRECO_E_FORMA_DE_PAGAMENTO]"),
    ("DA TRANSMISSÃO E DA POSSE", "[PREENCHER:TRANSMISSAO_E_POSSE]"),
    ("DAS DECLARAÇÕES DAS PARTES VENDEDORAS", "[PREENCHER:DECLARACOES_VENDEDORES]"),
    ("DOS TRIBUTOS", "[PREENCHER:TRIBUTOS_E_ITBI]"),
    ("DAS CERTIDÕES E CONSULTAS", "[PREENCHER:CERTIDOES_E_CONSULTAS]"),
    ("DAS DECLARAÇÕES DAS PARTES COMPRADORAS", "[PREENCHER:DECLARACOES_COMPRADORES]"),
    ("DOS ARQUIVAMENTOS", "[PREENCHER:ARQUIVAMENTOS]"),
    ("DO ENCERRAMENTO", "[PREENCHER:ENCERRAMENTO_DO_ATO]"),
]


def configurar_fonte(run, *, tamanho: float, negrito: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(tamanho)
    run.font.bold = negrito
    run.font.color.rgb = RGBColor(0, 0, 0)


def criar() -> None:
    documento = Document()
    documento.core_properties.title = "Modelo estruturado de escritura pública de venda e compra"
    documento.core_properties.subject = "Modelo reutilizável com blocos semânticos"
    documento.core_properties.author = "ThiagoDocs"
    documento.core_properties.last_modified_by = "ThiagoDocs"

    secao = documento.sections[0]
    secao.start_type = WD_SECTION.NEW_PAGE
    secao.page_width = Cm(21)
    secao.page_height = Cm(29.7)
    secao.top_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(2.5)
    secao.header_distance = Cm(1.25)
    secao.footer_distance = Cm(1.25)

    normal = documento.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(14)
    titulo.paragraph_format.keep_with_next = True
    configurar_fonte(
        titulo.add_run("ESCRITURA PÚBLICA DE VENDA E COMPRA"),
        tamanho=13,
        negrito=True,
    )

    for rotulo, marcador in SECOES:
        if rotulo:
            cabecalho = documento.add_paragraph()
            cabecalho.paragraph_format.space_before = Pt(8)
            cabecalho.paragraph_format.space_after = Pt(3)
            cabecalho.paragraph_format.keep_with_next = True
            configurar_fonte(
                cabecalho.add_run(rotulo),
                tamanho=10,
                negrito=True,
            )
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.paragraph_format.space_after = Pt(6)
        paragrafo.paragraph_format.line_spacing = 1.15
        configurar_fonte(paragrafo.add_run(marcador), tamanho=11)

    DESTINO.parent.mkdir(parents=True, exist_ok=True)
    documento.save(DESTINO)


if __name__ == "__main__":
    criar()
